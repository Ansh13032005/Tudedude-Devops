from docx import Document
from docx.shared import Pt, Inches

doc = Document()
doc.add_heading('DevOps Assignment Submission', 0)

p = doc.add_paragraph()
p.add_run('Name: ').bold = True
p.add_run('[Your Name]\n')
p.add_run('GitHub Repository Link: ').bold = True
p.add_run('[Your Repo Link Here]\n')

# Part 1
doc.add_heading('Part 1: GitHub Repository Setup & Flask Project', level=2)
p = doc.add_paragraph()
p.add_run('Goal: ').bold = True
p.add_run('Create a GitHub repository, clone it, create a branch named Tutedude, add a Flask project, and merge it into main.')
p = doc.add_paragraph()
p.add_run('Explanation:\n').bold = True
p.add_run('I cloned the remote repository to my local machine, created a new branch called Tutedude, and added app.py containing a basic Flask application with a /api route. Finally, I committed the file and merged the Tutedude branch back into main.')
p = doc.add_paragraph()
p.add_run('[Insert Screenshot of Part 1 Terminal Output or Git Log Here]').bold = True
p.style = 'Heading 3'

# Part 2
doc.add_heading('Part 2: Branch Updates', level=2)
p = doc.add_paragraph()
p.add_run('Goal: ').bold = True
p.add_run('Create a new branch Tutedude_new, update the /api JSON content, and merge the branch into main.')
p = doc.add_paragraph()
p.add_run('Explanation:\n').bold = True
p.add_run('I created a branch called Tutedude_new and modified the JSON response in the /api route. I then staged, committed the changes, and merged them back into the main branch.')
p = doc.add_paragraph()
p.add_run('[Insert Screenshot of Part 2 Terminal Output or Git Log Here]').bold = True
p.style = 'Heading 3'

# Part 3
doc.add_heading('Part 3: Parallel Feature Development (master_1 & master_2)', level=2)
p = doc.add_paragraph()
p.add_run('Goal: ').bold = True
p.add_run('Create branches master_1 (frontend) and master_2 (backend), perform parallel development, and merge both into main while resolving conflicts.')
p = doc.add_paragraph()
p.add_run('Explanation:\n').bold = True
p.add_run('I branched master_1 and master_2 simultaneously from main. In master_1, I created a To-Do HTML form. In master_2, I added a /submittodoitem backend route connected to MongoDB. When merging both into main, I resolved a merge conflict in app.py manually.')
p = doc.add_paragraph()
p.add_run('[Insert Screenshot of Part 3 Terminal Output or Git Log Here]').bold = True
p.style = 'Heading 3'

# Part 4
doc.add_heading('Part 4: Form Enhancements, Git Reset, and Rebase', level=2)
p = doc.add_paragraph()
p.add_run('Goal: ').bold = True
p.add_run('Add ID, UUID, and Hash fields to the To-Do form sequentially in master_1, merge to main, perform a soft reset in main, re-commit, and rebase main onto master_1.')
p = doc.add_paragraph()
p.add_run('Explanation:\n').bold = True
p.add_run('In master_1, I added the Item ID, Item UUID, and Item Hash fields to the To-Do form, making separate commits. After merging master_1 into main, I performed a git reset --soft back to the Item ID commit, re-committed the remaining changes as a single commit, and finally used git rebase to sync the rewritten history back into master_1.')
p = doc.add_paragraph()
p.add_run('[Insert Screenshot of Part 4 Terminal Output or Git Log Here]').bold = True
p.style = 'Heading 3'

# Final Push
doc.add_heading('Final Push', level=3)
doc.add_paragraph('All branches and commit histories were pushed to the remote GitHub repository.')

doc.save('submission.docx')
