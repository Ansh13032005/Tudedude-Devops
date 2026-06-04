import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right, insideH, insideV
    value is a dictionary like: {'sz': 12, 'val': 'single', 'color': 'FF0000', 'space': '0'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def create_report():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LAB REPORT ASSIGNMENT")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Flask & MongoDB Atlas Integration")
    sub_run.font.size = Pt(16)
    sub_run.italic = True
    
    doc.add_paragraph("\n" * 6)
    
    meta_info = doc.add_paragraph()
    meta_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_info.add_run("Submitted By: [Your Name]\n").bold = True
    meta_info.add_run("Course: DevOps / Full Stack Development\n")
    meta_info.add_run("Date: June 2026\n")
    
    doc.add_page_break()
    
    # Section 1: Overview
    doc.add_heading("1. Project Overview & Structure", level=1)
    p_intro = doc.add_paragraph(
        "This project implements a web application using Flask connected to a MongoDB Atlas cloud database. "
        "It accomplishes two main tasks: (1) serving a local JSON data source via a REST API endpoint, "
        "and (2) rendering a web form that processes submissions and persists them securely to a remote MongoDB Atlas cluster."
    )
    
    doc.add_heading("Project Structure Layout", level=2)
    struct_block = doc.add_paragraph()
    struct_run = struct_block.add_run(
        "flask-mongodb-assignment/\n"
        "├── app.py                     # Backend server and routing\n"
        "├── data.json                  # Mock API storage\n"
        "├── requirements.txt           # Dependency declaration\n"
        "├── templates/\n"
        "│   ├── index.html             # Main submission form\n"
        "│   └── success.html           # Submission confirmation\n"
        "├── static/                    # Directory for static files\n"
        "└── README.md                  # Project documentation\n"
    )
    struct_run.font.name = 'Courier New'
    struct_run.font.size = Pt(9.5)
    
    doc.add_heading("Step 1: Installation & Setup", level=2)
    doc.add_paragraph("Command executed to install dependencies:")
    cmd_p = doc.add_paragraph()
    cmd_run = cmd_p.add_run("pip install flask pymongo dnspython")
    cmd_run.font.name = 'Courier New'
    cmd_run.font.bold = True
    
    # Screenshot Placeholder Table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell, 
                    top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    cell.paragraphs[0].text = "[PLACEHOLDER: Attach Package Installation Screenshot Here]"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer
    
    doc.add_page_break()
    
    # Question 1
    doc.add_heading("2. Question 1: Serving REST API Data", level=1)
    
    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Create a '/api' endpoint in the Flask application that reads contents from a local data.json file and outputs a JSON response.")
    
    doc.add_heading("Backend Data (data.json)", level=2)
    json_p = doc.add_paragraph()
    json_run = json_p.add_run(
        "[\n"
        "    {\n"
        "        \"name\": \"Ansh\",\n"
        "        \"course\": \"DevOps\"\n"
        "    },\n"
        "    {\n"
        "        \"name\": \"Rahul\",\n"
        "        \"course\": \"Python\"\n"
        "    }\n"
        "]"
    )
    json_run.font.name = 'Courier New'
    json_run.font.size = Pt(10)
    
    doc.add_heading("Route Implementation (app.py snippet)", level=2)
    route1_p = doc.add_paragraph()
    route1_run = route1_p.add_run(
        "@app.route('/api')\n"
        "def get_api_data():\n"
        "    with open('data.json', 'r') as file:\n"
        "        data = json.load(file)\n"
        "    return jsonify(data)"
    )
    route1_run.font.name = 'Courier New'
    route1_run.font.size = Pt(10)
    
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(
        "When a client sends a GET request to the `/api` path, the server opens `data.json` in read-only mode, "
        "parses the JSON array using the standard library `json` module, and serializes it using Flask's `jsonify` function. "
        "This ensures proper `application/json` headers are returned."
    )
    
    doc.add_heading("Output Screenshot", level=2)
    table_api = doc.add_table(rows=1, cols=1)
    table_api.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_api = table_api.cell(0, 0)
    set_cell_border(cell_api, 
                    top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'},
                    right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    cell_api.paragraphs[0].text = "[PLACEHOLDER: Attach /api Response Screenshot Here]"
    cell_api.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer
    
    doc.add_heading("Result", level=2)
    doc.add_paragraph("Successfully retrieved and displayed JSON mock list of students.")
    
    doc.add_page_break()
    
    # Question 2
    doc.add_heading("3. Question 2: Student Registration Form & MongoDB Atlas Connection", level=1)
    
    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Design a web form using HTML that accepts a student's Name and Email. Submit the form data to a POST handler in Flask that stores it dynamically into a remote MongoDB Atlas database cluster.")
    
    doc.add_heading("HTML Form Template (templates/index.html)", level=2)
    html_p = doc.add_paragraph()
    html_run = html_p.add_run(
        "<!-- index.html excerpt -->\n"
        "<form action=\"/submit\" method=\"POST\">\n"
        "    <label>Name:</label>\n"
        "    <input type=\"text\" name=\"name\" required>\n"
        "    <label>Email:</label>\n"
        "    <input type=\"email\" name=\"email\" required>\n"
        "    <button type=\"submit\">Submit</button>\n"
        "</form>"
    )
    html_run.font.name = 'Courier New'
    html_run.font.size = Pt(9.5)
    
    doc.add_heading("Backend Handler (app.py snippet)", level=2)
    route2_p = doc.add_paragraph()
    route2_run = route2_p.add_run(
        "@app.route('/submit', methods=['POST'])\n"
        "def submit():\n"
        "    try:\n"
        "        name = request.form['name']\n"
        "        email = request.form['email']\n"
        "        data = {\"name\": name, \"email\": email}\n"
        "        collection.insert_one(data)\n"
        "        return redirect(url_for('success'))\n"
        "    except Exception as e:\n"
        "        return render_template('index.html', error=str(e))"
    )
    route2_run.font.name = 'Courier New'
    route2_run.font.size = Pt(10)
    
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(
        "Upon submission, the `/submit` route intercepts the POST payload. It extracts the `name` and `email` fields. "
        "The application constructs a dictionary and passes it to the MongoDB driver's `insert_one()` method. "
        "If successful, the client is redirected to `/success`. If an error occurs (such as database disconnection), "
        "the exception message is captured and rendered back onto the home page form in red text."
    )
    
    doc.add_heading("MongoDB Atlas Setup & Connection String", level=2)
    doc.add_paragraph(
        "A cluster named Cluster0 was deployed on AWS. A database user with read/write privileges was created, and network whitelist rules "
        "were adjusted to 0.0.0.0/0. The application connects using the PyMongo MongoClient:"
    )
    db_p = doc.add_paragraph()
    db_run = db_p.add_run("MONGO_URI = 'mongodb+srv://username:password@cluster.mongodb.net/?retryWrites=true&w=majority'")
    db_run.font.name = 'Courier New'
    db_run.font.size = Pt(9.5)
    
    doc.add_page_break()
    
    # Screenshots section for Form
    doc.add_heading("Verification & Outputs", level=2)
    
    doc.add_paragraph("1. Form Page View:")
    t_f = doc.add_table(rows=1, cols=1)
    t_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_f = t_f.cell(0, 0)
    set_cell_border(c_f, top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    c_f.paragraphs[0].text = "[PLACEHOLDER: Form Page Screen]"
    c_f.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph("2. Successful Submission confirmation:")
    t_s = doc.add_table(rows=1, cols=1)
    t_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_s = t_s.cell(0, 0)
    set_cell_border(c_s, top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    c_s.paragraphs[0].text = "[PLACEHOLDER: Success Page Screen]"
    c_s.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph("3. Database records confirmed in MongoDB Atlas:")
    t_m = doc.add_table(rows=1, cols=1)
    t_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_m = t_m.cell(0, 0)
    set_cell_border(c_m, top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    c_m.paragraphs[0].text = "[PLACEHOLDER: MongoDB Atlas Collections View Screen]"
    c_m.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph("4. Server logs showing terminal run:")
    t_t = doc.add_table(rows=1, cols=1)
    t_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_t = t_t.cell(0, 0)
    set_cell_border(c_t, top={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, bottom={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, left={'sz': 12, 'val': 'single', 'color': 'A0A0A0'}, right={'sz': 12, 'val': 'single', 'color': 'A0A0A0'})
    c_t.paragraphs[0].text = "[PLACEHOLDER: Running Server Terminal Screen]"
    c_t.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_heading("Result", level=2)
    doc.add_paragraph("Form data correctly transferred from client-side HTML to Python backend and successfully added as persistent JSON document to MongoDB cluster database.")
    
    # Save file to parent directory
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Flask_MongoDB_Assignment_Report.docx")
    doc.save(output_path)
    print(f"Report generated successfully at: {output_path}")

if __name__ == '__main__':
    create_report()
